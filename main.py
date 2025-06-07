import streamlit as st
import json
import os
import bleach
import re
from supabase import create_client, Client
from collections import defaultdict, Counter
from datetime import datetime
import time
import uuid
from streamlit_quill import st_quill

# Carregar variáveis de ambiente
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

st.set_page_config(page_title="Leitura de Leis por Cards", layout="centered")

# Configuração do Supabase
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
if not SUPABASE_URL or not SUPABASE_ANON_KEY or not SUPABASE_SERVICE_KEY:
    st.error("❌ Erro: As credenciais do Supabase (SUPABASE_URL, SUPABASE_ANON_KEY e SUPABASE_SERVICE_KEY) devem ser configuradas como variáveis de ambiente.")
    st.stop()
supabase: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)  # Para operações gerais
supabase_admin: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)  # Para operações administrativas como importação

# Função para carregar os dados do Supabase com paginação
def carregar_dados(usuario, start, end):
    response = supabase.table("cards").select("id, pergunta, resposta, referencia, concurso, lei, vezes_lido").eq("usuario", usuario).range(start, end).execute()
    dados = response.data if response.data else []
    return dados

# Função para contar o total de cards para o usuário
def contar_dados(usuario):
    response = supabase.table("cards").select("id", count="exact").eq("usuario", usuario).execute()
    return response.count if response.count else 0

# Função para contar o total de cards para uma lei específica
def contar_cards_por_lei(usuario, concurso, lei):
    response = supabase.table("cards").select("id", count="exact").eq("usuario", usuario).eq("concurso", concurso).eq("lei", lei).execute()
    return response.count if response.count else 0

# Função para carregar todas as leis para estatísticas e seletores
def carregar_leis(usuario):
    response = supabase.table("cards").select("lei").eq("usuario", usuario).execute()
    leis = sorted(set(item["lei"] for item in response.data if item.get("lei")))
    return leis

# Função para carregar estatísticas (para "Card mais lido por lei" e "Ranking de Leis Mais Lidas")
def carregar_estatisticas(usuario, leis_selecionadas=None):
    response = supabase.table("cards").select("lei, pergunta, vezes_lido").eq("usuario", usuario).execute()
    dados = response.data if response.data else []
    
    leituras_por_lei = Counter()
    mais_lido_por_lei = {}
    
    for item in dados:
        lei = item.get("lei", "[Sem Lei]")
        if leis_selecionadas and lei not in leis_selecionadas:
            continue
        leituras_por_lei[lei] += item.get("vezes_lido", 0)
        if lei not in mais_lido_por_lei or item.get("vezes_lido", 0) > mais_lido_por_lei[lei].get("vezes_lido", 0):
            mais_lido_por_lei[lei] = item
    
    mais_lidas = leituras_por_lei.most_common()
    return mais_lidas, mais_lido_por_lei

# Função para verificar se um card já existe (baseado na pergunta e resposta)
def card_existe(usuario, pergunta, resposta):
    response = supabase.table("cards").select("id").eq("usuario", usuario).eq("pergunta", pergunta).eq("resposta", resposta).execute()
    return len(response.data) > 0

# Função para salvar um card no Supabase
def salvar_card(usuario, card):
    data = {
        "usuario": usuario,
        "concurso": card["concurso"],
        "lei": card["lei"],
        "pergunta": card["pergunta"],
        "resposta": card["resposta"],
        "referencia": card["referencia"],
        "vezes_lido": card["vezes_lido"]
    }
    supabase_admin.table("cards").insert(data).execute()

# Função para atualizar um card no Supabase
def atualizar_card(usuario, card_antigo, card_novo):
    supabase.table("cards").update({
        "concurso": card_novo["concurso"],
        "lei": card_novo["lei"],
        "pergunta": card_novo["pergunta"],
        "resposta": card_novo["resposta"],
        "referencia": card_novo["referencia"],
        "vezes_lido": card_novo["vezes_lido"]
    }).eq("usuario", usuario).eq("pergunta", card_antigo["pergunta"]).eq("resposta", card_antigo["resposta"]).execute()

# Função para excluir um card do Supabase usando o id
def excluir_card(card_id):
    supabase.table("cards").delete().eq("id", card_id).execute()

# Função para carregar dados de um arquivo JSON (para importação)
def carregar_dados_json(arquivo):
    if isinstance(arquivo, str):
        with open(arquivo, "r", encoding="utf-8") as f:
            return json.load(f)
    else:
        return json.load(arquivo)

# Função para criar backup dos dados em formato JSON
def criar_backup(dados, usuario, session_id):
    if dados:
        os.makedirs("backup", exist_ok=True)
        nome_backup = f"backup/{usuario}_{session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(nome_backup, "w", encoding="utf-8") as f_backup:
            json.dump(dados, f_backup, ensure_ascii=False, indent=2)

# Função para validar o nome de usuário
def validar_usuario(usuario):
    usuario = usuario.strip().lower()
    if not re.match(r'^[a-z0-9_]+$', usuario):
        return None
    return usuario

# Função para exibir os cards filtrados e paginados
def exibir_cards(dados, total_cards, concurso_escolhido, lei_escolhida, fonte, usuario):
    # Ajustar o tamanho da fonte do título do expander via CSS sem interferir na animação
    st.markdown(
        f"""
        <style>
        div[data-testid="stExpander"] summary p {{
            font-size: {fonte}px !important;
        }}
        div[data-testid="stExpander"] {{
            transition: all 0.3s ease; /* Adicionar transição para animação */
        }}
        </style>
        """,
        unsafe_allow_html=True
    )

    filtro_leituras = st.selectbox(
        "Filtrar cards por número de leituras:",
        ["Todos", "Nunca lidos", "1 ou mais", "5 ou mais", "10 ou mais"]
    )

    busca = st.text_input("🔍 Buscar por palavra-chave, artigo, lei ou concurso:")

    # FILTRAGEM DOS CARDS
    perguntas_filtradas = []
    for i, item in enumerate(dados):
        vezes = item.get("vezes_lido", 0)
        concurso = item.get("concurso", "")
        lei = item.get("lei", "")
        if (
            concurso == concurso_escolhido and
            lei == lei_escolhida and
            (
                not busca or
                busca.lower() in item.get("pergunta", "").lower() or
                busca.lower() in item.get("resposta", "").lower() or
                busca.lower() in item.get("referencia", "").lower()
            ) and (
                filtro_leituras == "Todos" or
                (filtro_leituras == "Nunca lidos" and vezes == 0) or
                (filtro_leituras == "1 ou mais" and vezes >= 1) or
                (filtro_leituras == "5 ou mais" and vezes >= 5) or
                (filtro_leituras == "10 ou mais" and vezes >= 10)
            )
        ):
            perguntas_filtradas.append((i, item))

    # Contar o total de cards para a lei selecionada
    total_cards_lei = contar_cards_por_lei(usuario, concurso_escolhido, lei_escolhida)

    # Paginação
    PER_PAGE = 5
    total_paginas = max(1, (len(perguntas_filtradas) - 1) // PER_PAGE + 1)

    # Initialize or adjust pagina in session state
    if 'pagina' not in st.session_state or st.session_state['pagina'] > total_paginas:
        st.session_state['pagina'] = 1

    pagina_atual = st.sidebar.number_input(
        "Página", min_value=1, max_value=total_paginas, value=st.session_state['pagina'], step=1
    )
    st.session_state['pagina'] = pagina_atual

    inicio = (pagina_atual - 1) * PER_PAGE
    fim = min(inicio + PER_PAGE, len(perguntas_filtradas))

    # EXIBIÇÃO DOS CARDS
    if perguntas_filtradas:
        st.markdown(f"### 📑 Cards Cadastrados ({len(perguntas_filtradas)} de {total_cards_lei} cards)")

        for i, (index, item) in enumerate(perguntas_filtradas[inicio:fim]):
            pergunta_sanitizada = bleach.clean(
                item.get('pergunta', ''),
                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                strip=True
            )
            resposta_sanitizada = bleach.clean(
                item.get('resposta', ''),
                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                strip=True
            )

            pergunta_label = bleach.clean(item.get('pergunta', ''), tags=[], strip=True)

            with st.expander(f"📌 Pergunta (assunto): {pergunta_label}", expanded=False):
                st.markdown(f"<div style='font-size: {fonte}px;'><b>Resposta (conteúdo):</b> {resposta_sanitizada}</div>", unsafe_allow_html=True)
                st.caption(f"📖 Referência: {item.get('referencia', '')}  \n📘 Lei: {item.get('lei', '')}  \n🎯 Concurso: {item.get('concurso', '[Sem Concurso]')}")
                col1, col2, col3 = st.columns([1, 1, 1])

                with col1:
                    if st.button(f"✅ Lido ({item.get('vezes_lido', 0)}x)", key=f"btn_lido_{i}_{item.get('id', '')}"):
                        card_antigo = item.copy()
                        item["vezes_lido"] = item.get("vezes_lido", 0) + 1
                        atualizar_card(usuario, card_antigo, item)
                        st.session_state['pagina'] = pagina_atual
                        st.rerun()

                with col2:
                    if st.button("✏️ Editar", key=f"editar_{i}_{item.get('id', '')}"):
                        st.session_state["editar_id"] = item.get("id", "")  # Armazenar o ID do card

                with col3:
                    if st.button("🗑️ Excluir", key=f"excluir_{i}_{item.get('id', '')}"):
                        excluir_card(item.get("id", ""))
                        st.session_state['pagina'] = pagina_atual
                        st.rerun()

        # Botões de navegação entre páginas
        col_pag1, col_pag2 = st.columns(2)
        with col_pag1:
            if pagina_atual > 1 and st.button("⬅️ Página Anterior"):
                st.session_state['pagina'] = pagina_atual - 1
                st.rerun()
        with col_pag2:
            if pagina_atual < total_paginas and st.button("➡️ Próxima Página"):
                st.session_state['pagina'] = pagina_atual + 1
                st.rerun()

    else:
        st.info("ℹ️ Nenhum card encontrado com os filtros aplicados.")

    return perguntas_filtradas

# Inicializar estado de login e sessão
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
    st.session_state['usuario'] = None
if 'session_id' not in st.session_state:
    st.session_state['session_id'] = str(uuid.uuid4())

# 🔐 Login do usuário
if not st.session_state['logged_in']:
    usuario = st.text_input("🔐 Nome de usuário:")
    if usuario:
        usuario_valido = validar_usuario(usuario)
        if usuario_valido:
            st.session_state['logged_in'] = True
            st.session_state['usuario'] = usuario_valido
            st.rerun()
        else:
            st.error("❌ Nome de usuário inválido! Use apenas letras, números e sublinhados (ex.: joao123).")
    else:
        st.warning("Digite seu nome para continuar.")
    st.stop()

# Usuário logado
usuario = st.session_state['usuario']
session_id = st.session_state['session_id']

# Carregar o total de cards para paginação
total_cards = contar_dados(usuario)

# Carregar todas as leis para os seletores
leis_disponiveis = carregar_leis(usuario)

# Interface do Sidebar
st.sidebar.markdown("---")
fonte = st.sidebar.slider("🔠 Tamanho da Fonte (px):", 16, 48, 24)  # Aumentado o valor mínimo e padrão

# Botão de Logout
if st.sidebar.button("🚪 Sair"):
    st.session_state['logged_in'] = False
    st.session_state['usuario'] = None
    st.session_state['session_id'] = str(uuid.uuid4())
    st.session_state.clear()
    st.rerun()

st.markdown(f"<h1 style='font-size: {fonte + 20}px;'>📚 Leitura de Leis por Cards</h1>", unsafe_allow_html=True)
st.markdown(f"**Usuário logado:** {usuario}")

if 'leituras' not in st.session_state:
    st.session_state.leituras = {}

# Restaurar backup
st.sidebar.markdown("🛠️ **Restaurar Backup**")
if os.path.exists("backup"):
    arquivos_backup = sorted(
        [f for f in os.listdir("backup") if f.startswith(f"{usuario}_{session_id}")],
        reverse=True
    )
else:
    arquivos_backup = []

if arquivos_backup:
    escolha_backup = st.sidebar.selectbox("Selecione um backup para restaurar", arquivos_backup)
    if st.sidebar.button("♻️ Restaurar este backup"):
        caminho = os.path.join("backup", escolha_backup)
        dados_importados = carregar_dados_json(caminho)
        supabase.table("cards").delete().eq("usuario", usuario).execute()
        for item in dados_importados:
            if not card_existe(usuario, item["pergunta"], item["resposta"]):
                salvar_card(usuario, item)
        total_cards = contar_dados(usuario)  # Atualizar total_cards após restauração
        leis_disponiveis = carregar_leis(usuario)  # Atualizar leis_disponiveis
        st.session_state['pagina'] = 1
        st.rerun()
else:
    st.sidebar.caption("Nenhum backup encontrado.")

# Importar arquivo JSON
st.sidebar.markdown("📥 **Importar arquivo JSON personalizado**")
arquivo_json = st.sidebar.file_uploader("Escolha um arquivo .json", type="json")

if arquivo_json:
    if arquivo_json.size > 2 * 1024 * 1024:  # 2MB
        st.sidebar.error("❌ Arquivo muito grande! Limite: 2MB")
    elif st.sidebar.button("📂 Importar este arquivo"):
        criar_backup(carregar_dados(usuario, 0, total_cards - 1), usuario, session_id)
        dados_importados = carregar_dados_json(arquivo_json)
        novos_cards = 0
        duplicados = 0
        for item in dados_importados:
            if not card_existe(usuario, item["pergunta"], item["resposta"]):
                salvar_card(usuario, item)
                novos_cards += 1
                # Atualizar estado após cada inserção
                total_cards = contar_dados(usuario)
                leis_disponiveis = carregar_leis(usuario)
            else:
                duplicados += 1
        st.session_state['pagina'] = 1
        if duplicados > 0:
            st.sidebar.success(f"✅ {novos_cards} cards importados com sucesso! {duplicados} cards ignorados por já estarem cadastrados.")
        else:
            st.sidebar.success(f"✅ {novos_cards} cards importados com sucesso!")
        st.rerun()

st.markdown("## 🎯 Selecione um concurso para começar")

concursos_disponiveis = sorted(set(d["concurso"] for d in carregar_dados(usuario, 0, total_cards - 1) if d.get("concurso")))
concurso_escolhido = st.selectbox("Concurso:", ["Selecionar"] + concursos_disponiveis)

if concurso_escolhido != "Selecionar":
    leis_do_concurso = sorted(set(d["lei"] for d in carregar_dados(usuario, 0, total_cards - 1) if d.get("concurso") == concurso_escolhido and d.get("lei")))
    lei_escolhida = st.selectbox("📘 Lei do concurso:", ["Selecionar"] + leis_do_concurso)

    if lei_escolhida != "Selecionar":
        st.markdown(f"### Cards da Lei **{lei_escolhida}** para o Concurso **{concurso_escolhido}**")
        dados_completos = carregar_dados(usuario, 0, total_cards - 1)  # Carregar todos os dados
        perguntas_filtradas = exibir_cards(dados_completos, total_cards, concurso_escolhido, lei_escolhida, fonte, usuario)

        # ✏️ Editar Card
        if "editar_id" in st.session_state:
            card_id = st.session_state["editar_id"]
            st.markdown("---")
            st.subheader("✧️ Editar Card")
            response = supabase.table("cards").select("*").eq("id", card_id).eq("usuario", usuario).execute()
            item = response.data[0] if response.data else None

            if item:
                with st.form(f"form_editar_{card_id}"):
                    nova_pergunta = st_quill(
                        value=item["pergunta"],
                        placeholder="Digite a pergunta (assunto)...",
                        toolbar=["bold", "italic", "underline", "link", "list"],
                        html=True
                    )
                    nova_resposta = st_quill(
                        value=item["resposta"],
                        placeholder="Digite a resposta (conteúdo)...",
                        toolbar=["bold", "italic", "underline", "link", "list"],
                        html=True
                    )
                    nova_referencia = st.text_input("Referência", value=item["referencia"])
                    nova_concurso = st.text_input("Concurso", value=item["concurso"])
                    nova_lei = st.text_input("Lei", value=item["lei"])
                    confirmar = st.form_submit_button("Salvar alterações")

                    if confirmar:
                        if not nova_concurso or not nova_lei or not nova_pergunta or not nova_resposta:
                            st.error("❌ Todos os campos obrigatórios devem ser preenchidos!")
                        elif card_existe(usuario, nova_pergunta, nova_resposta) and (nova_pergunta != item["pergunta"] or nova_resposta != item["resposta"]):
                            st.error("❌ Um card com esta pergunta e resposta já existe!")
                        else:
                            nova_pergunta_sanitizada = bleach.clean(
                                nova_pergunta,
                                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                                strip=True
                            )
                            nova_resposta_sanitizada = bleach.clean(
                                nova_resposta,
                                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                                strip=True
                            )

                            novo_card = {
                                "usuario": usuario,
                                "pergunta": nova_pergunta_sanitizada,
                                "resposta": nova_resposta_sanitizada,
                                "referencia": nova_referencia,
                                "concurso": nova_concurso,
                                "lei": nova_lei,
                                "vezes_lido": item.get("vezes_lido", 0)
                            }
                            atualizar_card(usuario, item, novo_card)
                            del st.session_state["editar_id"]
                            st.session_state['pagina'] = 1
                            st.rerun()

# ESTATÍSTICAS
st.sidebar.markdown("---")
st.sidebar.markdown("📊 **Ranking de Leis Mais Lidas**")
leis_selecionadas_ranking = st.sidebar.multiselect(
    "Selecione as leis para o ranking:", leis_disponiveis, default=leis_disponiveis[:5] if len(leis_disponiveis) > 5 else leis_disponiveis
)
mais_lidas, _ = carregar_estatisticas(usuario, leis_selecionadas_ranking)
for lei, total in mais_lidas:
    st.sidebar.markdown(f"**{lei}** — {total} leituras")

st.sidebar.markdown("---")
st.sidebar.markdown("🔥 **Card mais lido por lei**")
leis_selecionadas_mais_lido = st.sidebar.multiselect(
    "Selecione as leis para os cards mais lidos:", leis_disponiveis, default=leis_disponiveis[:5] if len(leis_disponiveis) > 5 else leis_disponiveis
)
_, mais_lido_por_lei = carregar_estatisticas(usuario, leis_selecionadas_mais_lido)
for lei, item in mais_lido_por_lei.items():
    st.sidebar.markdown(f"**{lei}** → *{item['pergunta'][:50]}...* ({item['vezes_lido']}x)")

# Exportar para Word (Seletivo)
st.sidebar.markdown("📄 **Exportar para Word**")
export_concurso = st.sidebar.selectbox("Exportar cards do concurso:", ["Todos"] + concursos_disponiveis)
export_lei = st.sidebar.selectbox("Exportar cards da lei:", ["Todas"] + sorted(set(d["lei"] for d in carregar_dados(usuario, 0, total_cards - 1) if d.get("lei"))))

if st.sidebar.button("⬇️ Baixar cards selecionados em Word"):
    from docx import Document

    doc = Document()
    doc.add_heading("Cards de Estudo", 0)

    cards_filtrados = carregar_dados(usuario, 0, total_cards - 1)
    if export_concurso != "Todos":
        cards_filtrados = [d for d in cards_filtrados if d.get("concurso") == export_concurso]
    if export_lei != "Todas":
        cards_filtrados = [d for d in cards_filtrados if d.get("lei") == export_lei]

    if not cards_filtrados:
        st.sidebar.error("❌ Nenhum card encontrado com os filtros selecionados!")
    else:
        for item in cards_filtrados:
            doc.add_heading(item.get("concurso", "[Concurso]"), level=1)
            doc.add_heading(item.get("lei", "[Lei]"), level=2)
            doc.add_paragraph(f"Pergunta: {item['pergunta']}")
            doc.add_paragraph(f"Resposta: {item['resposta']}")
            doc.add_paragraph(f"Referência: {item['referencia']}")
            doc.add_paragraph("")

        caminho_word = f"cards_{usuario}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(caminho_word)
        with open(caminho_word, "rb") as file:
            st.download_button(
                label="📥 Clique aqui para baixar",
                data=file,
                file_name=caminho_word,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.sidebar.markdown("---")
st.sidebar.markdown("➕ **Cadastrar Novo Card**")

with st.sidebar.form("form_novo_card"):
    novo_concurso = st.text_input("Concurso")
    nova_lei = st.text_input("Lei")
    nova_pergunta = st_quill(
        placeholder="Digite a pergunta (assunto)...",
        toolbar=["bold", "italic", "underline", "link", "list"],
        html=True
    )
    nova_resposta = st_quill(
        placeholder="Digite a resposta (conteúdo)...",
        toolbar=["bold", "italic", "underline", "link", "list"],
        html=True
    )
    nova_referencia = st.text_input("Referência")
    st.caption("Use o editor para formatar o texto com negrito, itálico, sublinhado, links e listas.")
    cadastrar = st.form_submit_button("📌 Adicionar Card")

    if cadastrar:
        if not novo_concurso or not nova_lei or not nova_pergunta or not nova_resposta:
            st.sidebar.error("❌ Todos os campos obrigatórios devem ser preenchidos!")
        elif card_existe(usuario, nova_pergunta, nova_resposta):
            st.sidebar.error("❌ Um card com esta pergunta e resposta já existe!")
        else:
            nova_pergunta_sanitizada = bleach.clean(
                nova_pergunta,
                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                strip=True
            )
            nova_resposta_sanitizada = bleach.clean(
                nova_resposta,
                tags=['b', 'i', 'u', 'br', 'p', 'ul', 'ol', 'li', 'strong', 'em'],
                strip=True
            )

            novo_card = {
                "usuario": usuario,
                "concurso": novo_concurso,
                "lei": nova_lei,
                "pergunta": nova_pergunta_sanitizada,
                "resposta": nova_resposta_sanitizada,
                "referencia": nova_referencia,
                "vezes_lido": 0
            }
            salvar_card(usuario, novo_card)
            total_cards = contar_dados(usuario)  # Atualizar total_cards após cadastro
            leis_disponiveis = carregar_leis(usuario)  # Atualizar leis_disponiveis
            st.session_state['pagina'] = 1
            st.rerun()

st.sidebar.markdown("---")